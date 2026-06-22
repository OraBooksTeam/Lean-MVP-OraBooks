#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate OraBooks Stack Plan Word Document (Current Lean MVP + Future Enterprise).
Run: python generate_stack_plan_doc.py
"""

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "OraBooks_Stack_Plan_Current_and_Future.docx"

# Prefer Bangla-friendly fonts on Windows
FONT_BODY = "Nirmala UI"
FONT_FALLBACK = "Calibri"
ACCENT = RGBColor(0x1A, 0x56, 0x8E)


def set_run_font(run, size=11, bold=False, color=None, name=FONT_BODY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 22, 2: 16, 3: 13}.get(level, 11), bold=True, color=ACCENT)
    return h


def add_para(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=size)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=size)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = shading.get_or_add_shd()
        shd.set(qn("w:fill"), "1A568E")

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
                p.paragraph_format.space_after = Pt(2)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_diagram_block(doc, title, lines):
    add_para(doc, title, bold=True, size=11)
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    set_run_font(run, size=9, name="Consolas")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(10)


def build_document():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # ── Cover ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("OraBooks Stack Plan\nCurrent (Lean MVP) ও Future (Enterprise)")
    set_run_font(tr, size=24, bold=True, color=ACCENT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("TaxOra LLC · OraBooks Accounting SaaS\nBeginner-Friendly Technology Stack Report")
    set_run_font(sr, size=13)

    doc.add_paragraph()
    add_para(
        doc,
        "এই document-টি OraBooks Accounting SaaS-এর technology stack plan ব্যাখ্যা করে। "
        "প্রথমে দেখানো হয়েছে আজ Lean MVP-তে কী কী technology ব্যবহার হচ্ছে এবং কী security আছে। "
        "তারপর future Enterprise stack (NestJS, Next.js, PostgreSQL ইত্যাদি) কীভাবে integrate হবে তা বর্ণনা করা হয়েছে। "
        "শেষে Current ও Future stack-এর comparison table দেওয়া আছে।",
        size=11,
        space_after=12,
    )
    add_para(doc, "Document Version: 1.0  |  Date: June 2026  |  Project: Lean MVP OraBooks", size=10)
    doc.add_page_break()

    # ── Section 1: Current Stack ──
    add_heading(doc, "১. Current Stack Plan — Lean MVP (আজ যা আছে)", 1)
    add_para(
        doc,
        "OraBooks Lean MVP বর্তমানে WordPress Multisite-এর উপর একটি PHP plugin হিসেবে চলছে। "
        "UI-র জন্য React 19 + Vite + Tailwind 4 ব্যবহার করা হয়। "
        "Accounting data MySQL database-এ সংরক্ষিত হয়। "
        "SL-004 থেকে SL-139 পর্যন্ত specification অনুযায়ী module implement করা হয়েছে।",
    )

    add_heading(doc, "১.১ Current Technology Layers", 2)

    current_layers = [
        (
            "Platform — WordPress 6.x Multisite",
            "WordPress হলো hosting platform। Multisite দিয়ে subdomain-based multi-tenant routing করা যায়। "
            "Lean MVP-তে দ্রুত page, shortcode, cron এবং admin shell পাওয়ার জন্য এটি বেছে নেওয়া হয়েছে।",
        ),
        (
            "Backend — PHP 8.x Plugin",
            "সব accounting business logic PHP class-এ লেখা (includes/class-orabooks-*.php)। "
            "Journal posting, workflow, RBAC, partner commission—সব domain logic এখানে। "
            "WordPress hook-এর বাইরে plain PHP service pattern follow করা হয়েছে যাতে future-তে migrate করা সহজ হয়।",
        ),
        (
            "Frontend — React 19 + Vite 8 + Tailwind CSS 4",
            "orabooks-ui folder-এ React app build হয় এবং assets/react/ folder-এ WordPress plugin-এ enqueue হয়। "
            "Admin ও tenant frontend আলাদা bundle (admin.js, frontend.js)। "
            "Modern, fast UI Lean MVP-তে দ্রুত deliver করার জন্য।",
        ),
        (
            "Database — MySQL 8",
            "Accounting table গুলো {prefix}orabooks_* naming convention follow করে। "
            "Event table: {prefix}gob_*_tob (SL-302)। "
            "ACID transaction, FOR UPDATE lock (SL-301 workflow) এবং WordPress $wpdb integration-এর জন্য MySQL।",
        ),
        (
            "Events — SL-302 Transactional Outbox",
            "Business change-এর সাথে same transaction-এ event outbox-এ write হয়। "
            "Worker পরে consumer-দের deliver করে। Idempotency consumer_log দিয়ে ensure করা হয়। "
            "journal_posted, sale_delivered ইত্যাদি domain events support করে।",
        ),
        (
            "Queue — SL-303 Async Jobs + WordPress Cron",
            "Background job (OCR, export, webhook, email) orabooks_async_jobs table-এ queue হয়। "
            "WP Cron every_minute schedule দিয়ে worker process করে। "
            "MVP-তে Redis ছাড়াই async processing সম্ভব।",
        ),
        (
            "Workflow — SL-301 State Engine",
            "Journal, Invoice, Bill, Expense, Commission—সব record-এর status change একটিমাত্র entry point: "
            "OraBooks_Workflow::transition()। DB lock, audit, event publish built-in।",
        ),
        (
            "API — WordPress REST + Guarded AJAX + OpenAPI",
            "REST endpoint: /wp-json/api/*। OpenAPI spec: docs/openapi/openapi.json। "
            "JWT + org_id header দিয়ে tenant context pass হয়।",
        ),
        (
            "AI — Azure Document Intelligence + OpenAI/Azure OpenAI",
            "Receipt OCR (SL-028), Voice (SL-052), Classification (SL-022)। "
            "API key না থাকলে MVP stub fallback deterministic result দেয়।",
        ),
        (
            "PWA — Service Worker + IndexedDB",
            "Mobile expense receipt offline queue (SL-028)। Service worker shell cache করে। "
            "REST route: /wp-json/api/pwa/manifest ও service-worker।",
        ),
        (
            "Testing — PHPUnit 11 + Jest",
            "PHP unit test প্রতিটি SL module cover করে (tests/OraBooks_*_Test.php)। "
            "JavaScript test Jest + jsdom (tests/js/)।",
        ),
        (
            "Deploy — Manual PowerShell Build",
            "orabooks-ui/build-live.ps1 → npm run build → assets/react/। "
            "OraBooks_DeployChecks production gate verify করে (secrets, TLS, schema, cron)।",
        ),
        (
            "Cloud (Current) — Self-hosted / Shared Hosting",
            "MVP stage-এ low cost hosting (FTP-style deploy)। Docker/CI/CD এখনো নেই।",
        ),
    ]

    for title, body in current_layers:
        add_heading(doc, title, 3)
        add_para(doc, body)

    add_heading(doc, "১.২ Current Architecture (Diagram)", 2)
    add_diagram_block(
        doc,
        "Data Flow — Lean MVP:",
        [
            "  [Browser: React 19 + Vite + Tailwind UI]",
            "           |",
            "           v",
            "  [WordPress Multisite — Pages, Shortcodes, REST/AJAX]",
            "           |",
            "           v",
            "  [PHP Plugin — Posting, Workflow, RBAC, Partner, AI]",
            "       |         |              |",
            "       v         v              v",
            "  [MySQL]  [SL-302 Outbox]  [SL-303 Jobs + WP Cron]",
            "  orabooks_*     gob_*_tob       async_jobs",
        ],
    )

    add_heading(doc, "১.৩ Domain Modules (SL-004 → SL-139)", 2)
    add_bullets(
        doc,
        [
            "SL-004: Organization, subdomain, tier, data residency, partner/customer org types",
            "SL-013: Auth — register, login, Google OIDC, JWT, refresh token, email verify",
            "SL-003: RBAC/ABAC — deny-by-default permissions, cross-tenant guard",
            "SL-014: Team invites, role assignment",
            "SL-001: Journal posting engine — canonical ledger, hash chain",
            "SL-002: Approval gate — journal review workflow",
            "SL-301: Workflow state engine — invoice, bill, expense, commission states",
            "SL-017: Chart of Accounts",
            "SL-021: Customers & Invoices (AR)",
            "SL-027: Vendors & Bills (AP)",
            "SL-028: Expenses OCR + PWA offline receipts",
            "SL-031: Bank reconciliation",
            "SL-034: Inventory lite",
            "SL-052: Voice-to-text input",
            "SL-022: Smart classification + tax hints",
            "SL-076: AI review queue",
            "SL-068: Partner commission engine",
            "SL-139: Partner dashboard & onboarding",
            "SL-074/075: Financial & operational reports",
            "SL-114: PDF/CSV exports",
            "SL-113: CSV imports",
            "SL-203: Attachments & versioning",
            "SL-250: Notification center (email, push, in-app)",
            "SL-304: Fiscal period governance",
            "SL-305: Tax engine",
            "SL-302: Domain event bus",
            "SL-303: Async job queue",
            "SL-093: Observability & monitoring",
            "SL-008: Secrets & TLS",
            "SL-009: Audit logging",
            "SL-099: OWASP security controls",
        ],
    )

    doc.add_page_break()

    # ── Section 2: Current Security ──
    add_heading(doc, "২. Current Security Plan (আজ যা Implement করা আছে)", 1)
    add_para(
        doc,
        "OraBooks Lean MVP-তে security specification-driven design follow করা হয়েছে। "
        "নিচে প্রতিটি security layer বর্ণনা করা হয়েছে—কী আছে এবং কেন গুরুত্বপূর্ণ।",
    )

    security_items = [
        (
            "SL-008 — Secrets & TLS Management",
            [
                "Secret storage: Environment variable (ORABOOKS_*), JSON file, বা encrypted WordPress options",
                "AES-256-CBC encryption at rest — 2FA secret, backup codes, sensitive data",
                "JWT HS256 signing — 15 minute access token, grace period secret rotation (24h)",
                "bcrypt password hashing (cost 10)",
                "Production HTTPS redirect — HTTP request automatically HTTPS-এ redirect",
                "Database TLS detection — MYSQL_SSL, DB_SSL env check",
                "TLS certificate expiry monitoring — 30 day warning, expired = critical alert",
                "Secret redaction in logs — mask_value(), redact_sensitive()",
                "TOTP (RFC 6238) — authenticator app support",
                "Google OIDC id_token verification — RS256 + JWKS",
            ],
            "Accounting software-এ secret leak মানে financial data compromise। SL-008 সব key encrypt করে এবং production-এ TLS enforce করে।",
        ),
        (
            "SL-013 — Authentication & Session Security",
            [
                "Register/Login — email validation, subdomain binding, rate limit (5/hr registration, 5/15min login failure)",
                "Refresh token rotation — SHA-256 hashed in DB, single-use, device metadata (IP, User-Agent)",
                "Email verification — JWT issue হয় না verify না হওয়া পর্যন্ত",
                "Password reset — 1 hour token, reset-এ সব refresh token revoke",
                "Google OIDC — OAuth state CSRF protection (10 min transient)",
                "2FA challenge flow — 5 min JWT with purpose=2fa_challenge",
                "Partner accounting isolation — partner org accounting page/API block (403)",
                "Cross-subdomain session — one-time JWT exchange",
                "HTTP-only secure cookies — JWT + refresh token cookie",
            ],
            "Unauthorized access accounting data-এ direct financial loss। Multi-layer auth (password + 2FA + OIDC) tenant isolation ensure করে।",
        ),
        (
            "SL-013 — Two-Factor Authentication (2FA)",
            [
                "TOTP setup — QR code, provisioning URI, encrypted secret storage",
                "Backup codes — bcrypt hashed, OTP-gated reveal/regenerate",
                "Org-wide 2FA policy — require_2fa config, non-compliant user API block",
                "Admin 2FA recovery — platform admin বা org owner recovery with audit log",
                "Disable 2FA — OTP required, org mandate হলে block",
            ],
            "Password leak হলেও 2FA second barrier দেয়। Enterprise customer-দের জন্য org-wide 2FA mandatory করা যায়।",
        ),
        (
            "SL-003 — RBAC / ABAC Access Control",
            [
                "Deny-by-default permission matrix — ~30 permissions, owner/admin/approver/staff/viewer roles",
                "require_permission() middleware — org-scoped, active org validation",
                "Cross-tenant guard — target_org_id mismatch = block + audit",
                "Partner commission ABAC — optional staff/viewer commission access",
                "Permission denied audit — permission_audit_log table",
                "Role change logging — who changed whose role, when",
            ],
            "Multi-tenant SaaS-এ এক org-এর user অন্য org-এর data দেখতে পারবে না—এটা RBAC/ABAC enforce করে।",
        ),
        (
            "SL-009 — Audit Logging",
            [
                "Immutable audit_logs table — severity, IP, User-Agent, correlation_id",
                "Metadata sanitization — secret redaction, email masking/hashing",
                "CSV export with masked PII",
                "365-day retention → archive table",
                "Self-audit — audit log view/export itself logged",
                "Correlation ID — request-scoped tracing across workflow transitions",
            ],
            "Accounting-এ 'কে কখন কী করল' prove করতে audit log mandatory। Compliance ও dispute resolution-এ critical।",
        ),
        (
            "SL-099 — OWASP Top-10 Security Controls",
            [
                "Security headers — HSTS, CSP, X-Frame-Options, nosniff, Referrer-Policy, Permissions-Policy",
                "Centralized rate-limit config — registration, login, API, export limits",
                "SSRF allowlist — outbound URL validation, private/local host block",
                "Input schema validation — UUID, email, org_id, amount patterns",
                "Security incident tracking — security_incidents table (SSRF, rate limit, access denied)",
                "Access-denied burst alerting — >10 403/hour → critical notification",
                "Weekly dependency scan — Composer/npm inventory, vulnerability alert",
                "Ledger integrity check — hash chain validation cron",
                "Secret rotation reminder — 90-day policy, admin notification",
            ],
            "OWASP standard follow করে common web attack (XSS, injection, SSRF, brute force) prevent করা হয়।",
        ),
        (
            "SL-301 — Workflow Safety",
            [
                "Single write path — status field সরাসরি update করা যায় না, শুধু transition()",
                "DB transaction + SELECT FOR UPDATE lock — race condition prevent",
                "Fail-closed — invalid transition = 409 error + audit",
                "RBAC/fiscal preconditions — transition-এর আগে permission ও period check",
            ],
            "Concurrent request-এ duplicate posting বা invalid state change prevent করে—ledger integrity রক্ষা।",
        ),
        (
            "SL-304 — Fiscal Period Guards",
            [
                "OPEN / SOFT_CLOSED / HARD_CLOSED period states",
                "Closed period-এ posting block",
                "Workflow precondition integration",
            ],
            "Month-end close-এর পর accidental posting block—financial reporting accuracy।",
        ),
        (
            "Deploy Checks — Production Gates",
            [
                "JWT secret present and valid length",
                "Encryption key present",
                "HTTPS enabled in production",
                "Database TLS optional check",
                "Required tables exist (users, orgs, journals, async_jobs, etc.)",
                "Cron schedules registered",
                "Async job handlers registered",
            ],
            "Deploy-এর পর automatic verify—production-এ misconfiguration catch করে।",
        ),
    ]

    for title, bullets, why in security_items:
        add_heading(doc, title, 2)
        add_bullets(doc, bullets)
        add_para(doc, f"কেন গুরুত্বপূর্ণ: {why}", bold=False, size=10)

    doc.add_page_break()

    # ── Section 3: Future Stack ──
    add_heading(doc, "৩. Future Stack Plan — Enterprise (Expanded Edition)", 1)
    add_para(
        doc,
        "Lean MVP stable হওয়ার পর OraBooks Enterprise SaaS-এ upgrade হবে। "
        "নিচের technology stack modern, scalable এবং AI/automation-ready। "
        "প্রতিটি layer-এর Bangla ব্যাখ্যা সহ purpose দেওয়া হয়েছে।",
    )

    future_layers = [
        (
            "Frontend — Next.js + TypeScript",
            "Next.js-এর ভিতরেই React থাকে। SEO, Server Components, routing, PWA support built-in। "
            "Enterprise SaaS UI-এর industry standard। TypeScript type safety দিয়ে bug কমায়।",
        ),
        (
            "UI Design — Tailwind CSS + shadcn/ui",
            "Professional, accessible component library। Dashboard, table, form, dialog—সব consistent design system। "
            "Fast development with reusable components।",
        ),
        (
            "Backend — NestJS",
            "OraBooks শুধু accounting app নয়—RBAC, Audit, Inventory, AI, Approval Workflow, Event Bus, "
            "Multi-Tenant Architecture, API Layer, Automation সব আছে। "
            "NestJS structured module architecture দেয়—বড় SaaS maintain করা সহজ।",
        ),
        (
            "ORM — Prisma",
            "Object Relational Mapper—SQL হাতে না লিখে TypeScript code দিয়ে database access। "
            "Type safety, auto migration, fewer bugs, faster development।",
        ),
        (
            "Database — PostgreSQL",
            "Customers, Invoices, Journals, Ledger, Inventory, Users, Roles—সব data এখানে। "
            "Accounting-এ transaction safety critical—PostgreSQL ACID capability best choice।",
        ),
        (
            "AI Database — pgvector",
            "PostgreSQL extension for vector search। "
            "Future: AI Copilot, Smart Search, RAG, AI Knowledge Base, Natural Language Query।",
        ),
        (
            "Cache — Redis",
            "Ultra-fast memory database। Session, OTP, Permission Cache, Dashboard Cache, Rate Limiting, Queue backend। "
            "PostgreSQL load কমায়, application fast করে।",
        ),
        (
            "Queue — BullMQ",
            "Background Job Processing। Email, OCR, AI Analysis, Notification, Report Generation—"
            "user request block না করে background-এ process। UI smooth থাকে।",
        ),
        (
            "Storage — Cloudflare R2",
            "Receipt, PDF, Contract, Attachment, Invoice Copy, Document Versioning। "
            "File database-এ রাখা উচিত নয়। R2 low cost scalable storage + CDN support।",
        ),
        (
            "Search — Meilisearch",
            "Fast Search Engine। Customer, Invoice, Vendor, Product—instant search experience। "
            "PostgreSQL full-text search-এর চেয়ে অনেক faster।",
        ),
        (
            "AI Layer — OpenAI, Gemini, Claude",
            "Smart Classification, OCR Review, AI Approval Suggestions, Tax Hints, Expense Analysis, AI Copilot। "
            "Abstraction layer দিয়ে provider change করা easy।",
        ),
        (
            "Automation — n8n",
            "No-Code/Low-Code Workflow Automation। Invoice Paid → WhatsApp, Approval → Email, CRM Sync, "
            "AI Trigger, Google Sheet Integration।",
        ),
        (
            "Monitoring — Grafana + Prometheus + Loki",
            "Grafana dashboard, Prometheus metrics, Loki logs। "
            "Production-এ error, slow request, CPU, memory monitor।",
        ),
        (
            "Tracing — OpenTelemetry",
            "Request system-এর ভেতরে কোথায় কত সময় নিচ্ছে track করে। "
            "Example: Invoice save 5 second—OpenTelemetry বলবে কোন service slow।",
        ),
        (
            "Deployment — Docker",
            "Application container-এ run। Developer laptop ও production server same environment। "
            "'আমার PC-তে কাজ করে server-এ না'—এই problem solve।",
        ),
        (
            "CI/CD — GitHub Actions",
            "Code push → automatic test, build, deploy। Manual deployment কম, error কম।",
        ),
    ]

    for title, body in future_layers:
        add_heading(doc, title, 3)
        add_para(doc, body)

    add_heading(doc, "৩.১ Why NestJS?", 2)
    add_para(
        doc,
        "NestJS নির্বাচন করা হয়েছে কারণ OraBooks শুধুমাত্র Accounting App নয়। "
        "এতে RBAC, Audit, Inventory, CRM, AI, Approval Workflow, Event Bus, Multi-Tenant Architecture, "
        "API Layer এবং Automation থাকবে। NestJS বড় SaaS Application-এর জন্য structured architecture দেয় "
        "এবং code maintain করা সহজ করে।",
    )

    add_heading(doc, "৩.২ Why Next.js?", 2)
    add_para(
        doc,
        "Next.js-এর ভিতরেই React থাকে। তাই React আলাদা framework হিসেবে manage করতে হয় না। "
        "Next.js SEO, Performance, Routing, Server Components, API integration, PWA support এবং "
        "enterprise SaaS UI-এর জন্য industry standard।",
    )

    add_heading(doc, "৩.৩ Cloud Growth Path", 2)
    add_bullets(
        doc,
        [
            "MVP / Early Stage — Hetzner: কম খরচে ভালো performance, self-managed services",
            "1,000 → 10,000 Users — DigitalOcean: Managed PostgreSQL, Managed Redis, Kubernetes, Object Storage, Monitoring, better developer experience",
            "10,000+ Users — AWS: AI, OCR, Event Bus, Multi-Tenant, Audit, Compliance, Enterprise customers—AWS ecosystem best fit",
        ],
    )
    add_para(
        doc,
        "DigitalOcean example: 'Create Managed Database → Click → Done'। "
        "Hetzner-এ নিজে PostgreSQL চালাতে হয়, backup করতে হয়, monitor করতে হয়। "
        "DigitalOcean managed services developer time বাঁচায়।",
        size=10,
    )

    add_heading(doc, "৩.৪ Future Architecture (Diagram)", 2)
    add_diagram_block(
        doc,
        "Data Flow — Enterprise:",
        [
            "  [Browser: Next.js + TypeScript + Tailwind + shadcn/ui]",
            "           |",
            "           v",
            "  [NestJS Backend — Modules: Auth, RBAC, Posting, Workflow, AI]",
            "     |    |    |    |    |    |    |",
            "     v    v    v    v    v    v    v",
            "  [Prisma] [Redis] [BullMQ] [R2] [Meili] [AI] [n8n]",
            "     |",
            "     v",
            "  [PostgreSQL + pgvector]",
            "           |",
            "           v",
            "  [OpenTelemetry → Grafana / Prometheus / Loki]",
        ],
    )

    doc.add_page_break()

    # ── Section 4: Future Security ──
    add_heading(doc, "৪. Future Security Plan (Enterprise-তে আরও কী যোগ হবে)", 1)
    add_para(
        doc,
        "Current Lean MVP security solid foundation দিয়েছে। Enterprise stage-এ নিচের enhancement যোগ হবে।",
    )

    future_security_rows = [
        ("Authentication", "JWT + Google OIDC + 2FA", "Auth0/Clerk বা Keycloak; Enterprise SSO/SAML; Passkeys/WebAuthn"),
        ("Secrets Management", "Env + encrypted WP options", "HashiCorp Vault / AWS Secrets Manager; fully automated rotation"),
        ("Rate Limiting", "WordPress transient-based", "Redis-backed distributed rate limits across all servers"),
        ("WAF / DDoS", "Security headers (CSP, HSTS)", "Cloudflare WAF + AWS Shield Advanced"),
        ("Encryption", "AES-256-CBC plugin-level", "Field-level encryption; AWS KMS-managed keys"),
        ("Audit Log", "MySQL audit_logs table", "Immutable S3 + WORM storage; SIEM integration (Splunk/Datadog)"),
        ("Compliance", "OWASP self-check matrix", "SOC 2 Type II; GDPR tooling; regional data residency enforcement"),
        ("Network Security", "HTTPS + DB TLS", "Private VPC; mTLS between microservices"),
        ("Dependency Security", "Weekly inventory cron", "Snyk/Dependabot in GitHub Actions; auto-fix PRs"),
        ("Penetration Testing", "Not in MVP scope", "Annual third-party pentest; bug bounty program"),
        ("Backup & Recovery", "Hosting-provider dependent", "PostgreSQL PITR; R2 versioning; disaster recovery runbook"),
        ("Multi-Tenant Isolation", "org_id query scoping", "PostgreSQL Row-Level Security (RLS); automated tenant isolation tests"),
        ("AI Security", "Provider stubs + rate limits", "Prompt injection guards; PII redaction before AI; audit all AI decisions"),
        ("Session Management", "JWT + refresh cookie", "Centralized session store (Redis); device management dashboard"),
        ("Zero Trust", "Partner/customer ingress block", "Service mesh; per-request policy evaluation"),
    ]

    add_table(
        doc,
        ["Security Area", "Current (Lean MVP)", "Future (Enterprise)"],
        future_security_rows,
        col_widths=[1.4, 2.0, 2.4],
    )

    doc.add_page_break()

    # ── Section 5: Migration Roadmap ──
    add_heading(doc, "৫. Migration Roadmap (Current → Future)", 1)

    add_heading(doc, "Phase 1 — Lean MVP (এখন)", 2)
    add_bullets(
        doc,
        [
            "WordPress plugin continue — SL-001, SL-301, SL-302, SL-303 stabilize",
            "React UI inside WordPress pages",
            "MySQL database + WP Cron jobs",
            "Manual PowerShell deploy (build-live.ps1)",
            "Azure DI + OpenAI with MVP stubs",
            "Self-hosted / shared hosting",
        ],
    )

    add_heading(doc, "Phase 2 — Hybrid (৬–১২ মাস)", 2)
    add_bullets(
        doc,
        [
            "Strangler pattern: NestJS API layer extract — domain logic gradually move",
            "Next.js frontend replace React-in-WP pages one module at a time",
            "PostgreSQL + Prisma for new modules; MySQL → PostgreSQL migration scripts",
            "Redis + BullMQ replace high-frequency WP Cron workers",
            "Docker containers + GitHub Actions CI/CD pipeline",
            "Cloudflare R2 for new file uploads (receipts, exports, attachments)",
            "DigitalOcean managed PostgreSQL + Redis (1,000–10,000 users stage)",
        ],
    )

    add_heading(doc, "Phase 3 — Full Enterprise (১২–২৪ মাস)", 2)
    add_bullets(
        doc,
        [
            "WordPress retired or admin-only legacy shell",
            "Full NestJS backend + Next.js frontend + PostgreSQL",
            "Meilisearch instant search across all entities",
            "pgvector AI copilot + RAG knowledge base",
            "n8n workflow automation (WhatsApp, CRM, approval triggers)",
            "OpenTelemetry + Grafana/Prometheus/Loki full observability",
            "DigitalOcean → AWS migration when enterprise/compliance thresholds hit",
            "SOC 2 compliance program start",
        ],
    )

    add_heading(doc, "৫.১ What Stays the Same", 2)
    add_para(
        doc,
        "Technology stack change হলেও business rules same থাকবে: posting engine logic, workflow state machines, "
        "event types, RBAC permission matrix, fiscal period rules, tax calculation rules। "
        "শুধু runtime platform change—domain knowledge preserve হয়।",
    )

    doc.add_page_break()

    # ── Section 6: Stack Tables ──
    add_heading(doc, "৬. Stack Comparison Tables", 1)

    add_heading(doc, "Table A — Current Lean MVP Stack", 2)
    current_table_rows = [
        ("Frontend", "React 19 + Vite + Tailwind 4", "UI", "Fast MVP UI inside WordPress"),
        ("Backend", "PHP + WordPress Multisite", "Business Logic", "Rapid multi-tenant host + SL spec delivery"),
        ("Database", "MySQL 8", "Data Storage", "WP-native, FOR UPDATE workflow locks"),
        ("Events", "SL-302 Outbox (MySQL)", "Domain Events", "Transactional consistency"),
        ("Queue", "SL-303 + WP Cron", "Background Jobs", "MVP async without Redis infra"),
        ("Workflow", "SL-301 State Engine", "Status Management", "Single write path, audit, locks"),
        ("AI", "Azure DI + OpenAI", "Intelligence", "Spec-aligned OCR/voice/classification"),
        ("Storage", "Local / WP Uploads", "Files", "MVP simplicity"),
        ("Auth", "JWT + OIDC + 2FA", "Security", "SL-013 compliant"),
        ("API", "REST + AJAX + OpenAPI", "Integration", "WordPress-native API surface"),
        ("PWA", "Service Worker + IndexedDB", "Mobile/Offline", "Offline receipt queue"),
        ("Testing", "PHPUnit + Jest", "Quality", "SL test coverage per module"),
        ("Deploy", "Manual PowerShell", "Release", "Current FTP-style workflow"),
        ("Monitoring", "SL-093 Internal Metrics", "Observability", "MVP health dashboard"),
        ("Cloud", "Self-hosted / Shared", "Hosting", "Low cost MVP stage"),
    ]
    add_table(
        doc,
        ["Layer", "Technology", "Purpose", "Why Selected"],
        current_table_rows,
        col_widths=[1.0, 1.5, 1.2, 1.8],
    )

    add_heading(doc, "Table B — Future Enterprise Stack", 2)
    future_table_rows = [
        ("Frontend", "Next.js", "UI", "Modern SaaS frontend"),
        ("Language", "TypeScript", "Safety", "Type-safe code"),
        ("UI", "Tailwind + shadcn/ui", "Design", "Fast professional UI"),
        ("Backend", "NestJS", "Business Logic", "Enterprise architecture"),
        ("ORM", "Prisma", "Database Access", "Easy and safe DB operations"),
        ("Database", "PostgreSQL", "Data Storage", "ACID transactions"),
        ("AI DB", "pgvector", "Vector Search", "Future AI features"),
        ("Cache", "Redis", "Performance", "Fast caching"),
        ("Queue", "BullMQ", "Background Jobs", "Smooth UX"),
        ("Storage", "Cloudflare R2", "Files", "Scalable storage"),
        ("Search", "Meilisearch", "Search", "Instant search"),
        ("AI", "OpenAI / Gemini / Claude", "Intelligence", "AI features"),
        ("Automation", "n8n", "Workflow", "Process automation"),
        ("Monitoring", "Grafana / Prometheus / Loki", "Observability", "System health"),
        ("Tracing", "OpenTelemetry", "Tracing", "Performance debugging"),
        ("Deployment", "Docker", "Containers", "Portable apps"),
        ("CI/CD", "GitHub Actions", "Automation", "Auto deployment"),
        ("Cloud", "Hetzner → DigitalOcean → AWS", "Hosting", "Growth path"),
    ]
    add_table(
        doc,
        ["Layer", "Technology", "Purpose", "Why Selected"],
        future_table_rows,
        col_widths=[1.0, 1.5, 1.2, 1.8],
    )

    add_heading(doc, "৬.১ Quick Comparison Summary", 2)
    add_para(
        doc,
        "Lean MVP stack দ্রুত deliver এবং low cost-এর জন্য optimize। "
        "Future Enterprise stack scale, AI, automation এবং compliance-এর জন্য optimize। "
        "Migration gradual (strangler pattern)—business logic preserve, runtime upgrade।",
    )

    # Footer note
    doc.add_paragraph()
    add_para(
        doc,
        "— End of Document —\n"
        "OraBooks Stack Plan v1.0 | TaxOra LLC | Generated from Lean MVP codebase + Enterprise Stack specification",
        size=9,
    )

    return doc


def main():
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
